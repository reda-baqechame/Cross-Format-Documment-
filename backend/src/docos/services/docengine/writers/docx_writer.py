"""Universal canonical-model → DOCX writer.

Rebuilds a real ``.docx`` from the node graph regardless of the source format, so a
PDF- or TXT-origin document can be downloaded as Word. It is intentionally tolerant:
unknown node types are skipped and missing formatting is simply omitted rather than
raising, because fidelity-best-effort beats a hard failure on an unusual document.
"""

from __future__ import annotations

import io
import re
import zipfile
from xml.etree import ElementTree as ET

from docx import Document as DocxDocument
from docx.shared import Inches, Pt

from docos.model.document import CanonicalDocument
from docos.model.nodes import AnyNode
from docos.services.docengine.writers.redaction import is_redacted, node_text, run_text

# Cap embedded images to the printable width so a large picture never overflows the page.
_MAX_IMAGE_WIDTH = Inches(6)
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_W = f"{{{_W_NS}}}"
_R = f"{{{_R_NS}}}"
_REL = f"{{{_REL_NS}}}"
_CT = f"{{{_CT_NS}}}"
_FOOTNOTE_PLACEHOLDER = "DOCOS_FOOTNOTE_REF"

ET.register_namespace("w", _W_NS)
ET.register_namespace("r", _R_NS)
ET.register_namespace("", _CT_NS)


def model_to_docx(doc: CanonicalDocument, images: dict[str, bytes] | None = None) -> bytes:
    """Serialise ``doc`` to DOCX. ``images`` maps ``ImageNode.blob_ref`` → bytes; when an image's
    bytes are present they are embedded, otherwise a ``[image: …]`` placeholder is written."""
    images = images or {}
    out = DocxDocument()
    if doc.meta.title:
        out.core_properties.title = doc.meta.title

    pages_seen = 0
    for node in doc.children_of(doc.root_id):
        if node.type == "page":
            if pages_seen:
                out.add_page_break()
            pages_seen += 1
        _write_block(out, doc, node, images)

    buf = io.BytesIO()
    out.save(buf)
    data = buf.getvalue()
    return _inject_footnotes(data, doc) if _has_footnotes(doc) else data


def _write_block(
    out: DocxDocument, doc: CanonicalDocument, node: AnyNode, images: dict[str, bytes]
) -> None:
    kind = node.type

    if kind in ("root", "page"):
        for child in doc.children_of(node.id):
            _write_block(out, doc, child, images)

    elif kind == "heading":
        level = min(max(int(getattr(node, "level", 1) or 1), 0), 9)
        para = out.add_heading("", level=level)
        _add_runs(para, doc, node)

    elif kind == "paragraph":
        _add_runs(out.add_paragraph(), doc, node)

    elif kind == "list":
        ordered = bool(getattr(node, "ordered", False))
        style = "List Number" if ordered else "List Bullet"
        for item in doc.children_of(node.id):
            if item.type == "list_item":
                _add_runs(out.add_paragraph(style=style), doc, item)

    elif kind == "list_item":
        _add_runs(out.add_paragraph(style="List Bullet"), doc, node)

    elif kind == "table":
        _write_table(out, doc, node)

    elif kind == "image":
        if not is_redacted(doc, node.id):
            _write_image(out, doc, node, images)

    elif kind == "field":
        if not is_redacted(doc, node.id):
            name = getattr(node, "field_name", "field")
            out.add_paragraph(f"{name}: {node_text(doc, node)}")

    elif kind == "footnote":
        # Footnotes are emitted via the OOXML footnotes part after the main DOCX is built.
        return

    elif kind == "unsupported":
        label = getattr(node, "original_type", "unknown")
        out.add_paragraph(f"[unsupported node: {label}]")
        for child in doc.children_of(node.id):
            _write_block(out, doc, child, images)

    # comment / annotation / metadata_block / run are not emitted as standalone blocks.


def _write_image(
    out: DocxDocument, doc: CanonicalDocument, node: AnyNode, images: dict[str, bytes]
) -> None:
    """Embed the real image bytes when available, else fall back to a text placeholder."""
    data = images.get(getattr(node, "blob_ref", "") or "")
    if data:
        try:
            out.add_picture(io.BytesIO(data), width=_MAX_IMAGE_WIDTH)
            return
        except Exception:  # noqa: BLE001 - unreadable/unsupported image → placeholder
            pass
    out.add_paragraph(f"[image: {node_text(doc, node) or 'image'}]")


def _add_runs(para, doc: CanonicalDocument, block: AnyNode) -> None:
    for child in doc.children_of(block.id):
        if child.type == "footnote_reference":
            if not is_redacted(doc, child.id):
                para.add_run(f"[[{_FOOTNOTE_PLACEHOLDER}:{child.footnote_id}]]")
            continue
        if child.type == "unsupported":
            label = getattr(child, "original_type", "unknown")
            para.add_run(f"[unsupported: {label}]")
            continue
        if child.type != "run":
            continue
        text = run_text(doc, child)
        if not text:
            continue
        run = para.add_run(text)
        run.bold = bool(getattr(child, "bold", False))
        run.italic = bool(getattr(child, "italic", False))
        run.underline = bool(getattr(child, "underline", False))
        font_name = getattr(child, "font", None)
        if font_name:
            run.font.name = font_name
        size = getattr(child, "size", None)
        if size:
            run.font.size = Pt(float(size))


def _write_table(out: DocxDocument, doc: CanonicalDocument, tnode: AnyNode) -> None:
    rows = [n for n in doc.children_of(tnode.id) if n.type == "table_row"]
    if not rows:
        return
    ncols = max(
        (sum(1 for c in doc.children_of(r.id) if c.type == "table_cell") for r in rows),
        default=0,
    )
    if ncols == 0:
        return

    table = out.add_table(rows=len(rows), cols=ncols)
    for ri, row in enumerate(rows):
        cells = [c for c in doc.children_of(row.id) if c.type == "table_cell"]
        for ci, cell in enumerate(cells):
            if ci >= ncols:
                break
            text = _inline_text(doc, cell)
            table.cell(ri, ci).text = text


def _inline_text(doc: CanonicalDocument, block: AnyNode) -> str:
    parts: list[str] = []
    for child in doc.children_of(block.id):
        if child.type == "run":
            parts.append(run_text(doc, child))
        elif child.type == "footnote_reference" and not is_redacted(doc, child.id):
            parts.append(f"[{getattr(child, 'marker', '')}]")
        elif child.type == "unsupported":
            parts.append(f"[unsupported: {getattr(child, 'original_type', 'unknown')}]")
    return "".join(parts)


def _has_footnotes(doc: CanonicalDocument) -> bool:
    return any(n.type == "footnote_reference" for n in doc.nodes.values()) or any(
        n.type == "footnote" for n in doc.nodes.values()
    )


def _footnote_nodes(doc: CanonicalDocument) -> list[AnyNode]:
    return [n for n in doc.nodes.values() if n.type == "footnote"]


def _footnote_text_lines(doc: CanonicalDocument, footnote: AnyNode) -> list[str]:
    if is_redacted(doc, footnote.id):
        return [""]
    lines: list[str] = []
    direct = _inline_text(doc, footnote).strip()
    if direct:
        lines.append(direct)
    for child in doc.children_of(footnote.id):
        if child.type in ("paragraph", "heading", "table_cell"):
            text = _inline_text(doc, child).strip()
            if text:
                lines.append(text)
    return lines or [""]


def _footnote_id_map(doc: CanonicalDocument) -> dict[str, int]:
    ordered: list[str] = []
    for node in doc.walk():
        if node.type == "footnote_reference" and getattr(node, "footnote_id", None) not in ordered:
            ordered.append(node.footnote_id)
    for note in _footnote_nodes(doc):
        if note.footnote_id not in ordered:
            ordered.append(note.footnote_id)
    return {fid: i + 2 for i, fid in enumerate(ordered)}


def _inject_footnotes(data: bytes, doc: CanonicalDocument) -> bytes:
    id_map = _footnote_id_map(doc)
    if not id_map:
        return data

    with zipfile.ZipFile(io.BytesIO(data), "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}

    entries["word/document.xml"] = _replace_footnote_placeholders(
        entries["word/document.xml"], id_map
    )
    entries["word/footnotes.xml"] = _build_footnotes_xml(doc, id_map)
    entries["word/_rels/document.xml.rels"] = _ensure_footnotes_relationship(
        entries.get("word/_rels/document.xml.rels")
    )
    entries["[Content_Types].xml"] = _ensure_footnotes_content_type(entries["[Content_Types].xml"])

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, payload in entries.items():
            zout.writestr(name, payload)
    return out.getvalue()


def _replace_footnote_placeholders(document_xml: bytes, id_map: dict[str, int]) -> bytes:
    root = ET.fromstring(document_xml)
    pattern = re.compile(r"\[\[" + re.escape(_FOOTNOTE_PLACEHOLDER) + r":([^\]]+)\]\]")
    for run in root.findall(f".//{_W}r"):
        text_el = run.find(f"{_W}t")
        if text_el is None or not text_el.text:
            continue
        match = pattern.fullmatch(text_el.text)
        if not match:
            continue
        fid = match.group(1)
        docx_id = id_map.get(fid)
        if docx_id is None:
            continue
        for child in list(run):
            run.remove(child)
        rpr = ET.SubElement(run, f"{_W}rPr")
        ET.SubElement(rpr, f"{_W}rStyle", {f"{_W}val": "FootnoteReference"})
        ET.SubElement(run, f"{_W}footnoteReference", {f"{_W}id": str(docx_id)})
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _build_footnotes_xml(doc: CanonicalDocument, id_map: dict[str, int]) -> bytes:
    root = ET.Element(f"{_W}footnotes")
    sep = ET.SubElement(root, f"{_W}footnote", {f"{_W}type": "separator", f"{_W}id": "-1"})
    sep_p = ET.SubElement(sep, f"{_W}p")
    sep_r = ET.SubElement(sep_p, f"{_W}r")
    ET.SubElement(sep_r, f"{_W}separator")
    cont = ET.SubElement(
        root, f"{_W}footnote", {f"{_W}type": "continuationSeparator", f"{_W}id": "0"}
    )
    cont_p = ET.SubElement(cont, f"{_W}p")
    cont_r = ET.SubElement(cont_p, f"{_W}r")
    ET.SubElement(cont_r, f"{_W}continuationSeparator")

    notes_by_id = {note.footnote_id: note for note in _footnote_nodes(doc)}
    for fid, docx_id in id_map.items():
        note = notes_by_id.get(fid)
        footnote = ET.SubElement(root, f"{_W}footnote", {f"{_W}id": str(docx_id)})
        lines = _footnote_text_lines(doc, note) if note is not None else [""]
        for idx, line in enumerate(lines):
            p = ET.SubElement(footnote, f"{_W}p")
            ppr = ET.SubElement(p, f"{_W}pPr")
            ET.SubElement(ppr, f"{_W}pStyle", {f"{_W}val": "FootnoteText"})
            if idx == 0:
                marker = ET.SubElement(p, f"{_W}r")
                marker_pr = ET.SubElement(marker, f"{_W}rPr")
                ET.SubElement(marker_pr, f"{_W}rStyle", {f"{_W}val": "FootnoteReference"})
                ET.SubElement(marker, f"{_W}footnoteRef")
            if line:
                text_run = ET.SubElement(p, f"{_W}r")
                text = ET.SubElement(text_run, f"{_W}t")
                text.text = line
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _ensure_footnotes_relationship(rels_xml: bytes | None) -> bytes:
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    if rels_xml is None:
        root = ET.Element(f"{_REL}Relationships")
    else:
        root = ET.fromstring(rels_xml)
    for rel in root.findall(f"{_REL}Relationship"):
        if rel.get("Type") == rel_type:
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    existing = {rel.get("Id") for rel in root.findall(f"{_REL}Relationship")}
    rid = "rIdDocosFootnotes"
    if rid in existing:
        i = 1
        while f"{rid}{i}" in existing:
            i += 1
        rid = f"{rid}{i}"
    ET.SubElement(
        root,
        f"{_REL}Relationship",
        {"Id": rid, "Type": rel_type, "Target": "footnotes.xml"},
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _ensure_footnotes_content_type(content_types_xml: bytes) -> bytes:
    root = ET.fromstring(content_types_xml)
    part_name = "/word/footnotes.xml"
    for override in root.findall(f"{_CT}Override"):
        if override.get("PartName") == part_name:
            return ET.tostring(root, encoding="utf-8", xml_declaration=True)
    ET.SubElement(
        root,
        f"{_CT}Override",
        {
            "PartName": part_name,
            "ContentType": (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.footnotes+xml"
            ),
        },
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)
