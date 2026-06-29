"""DOCX adapter (python-docx) — the second functional adapter.

Extracts paragraphs (with heading detection), inline runs with basic formatting,
tables, and core metadata. Embedded metadata is preserved in ``meta.custom`` so the
provenance layer can flag and sanitize it.
"""

from __future__ import annotations

import io
import zipfile
from datetime import UTC, datetime
from xml.etree import ElementTree as ET

from docx import Document as DocxDocument

from docos.model.document import CanonicalDocument, DocumentMeta
from docos.model.ids import new_doc_id, new_node_id
from docos.model.nodes import (
    FootnoteNode,
    FootnoteReferenceNode,
    HeadingNode,
    ParagraphNode,
    RootNode,
    RunNode,
    TableCellNode,
    TableNode,
    TableRowNode,
)
from docos.services.docengine.interface import FormatAdapter
from docos.storage.blob import BlobStore

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_W_NS}}}"


class DocxAdapter(FormatAdapter):
    format_id = "docx"
    supported_mimes = (_DOCX_MIME,)

    def parse(self, data: bytes, *, blob: BlobStore | None = None) -> CanonicalDocument:
        src = DocxDocument(io.BytesIO(data))
        footnotes = _parse_footnotes(data)
        now = datetime.now(UTC)
        root = RootNode(id=new_node_id("root"))

        core = src.core_properties
        meta = DocumentMeta(
            title=core.title or None,
            author=core.author or None,
            source_format="docx",
            source_mime=_DOCX_MIME,
            created_at=now,
            modified_at=now,
            page_count=1,
            custom={
                "category": core.category,
                "comments": core.comments,
                "keywords": core.keywords,
                "last_modified_by": core.last_modified_by,
                "revision": core.revision,
                "subject": core.subject,
            },
        )

        doc = CanonicalDocument(doc_id=new_doc_id(), root_id=root.id, meta=meta)
        doc.add_node(root)

        order = 0
        for para in src.paragraphs:
            style = (para.style.name if para.style else "") or ""
            node: ParagraphNode | HeadingNode
            if style.startswith("Heading"):
                level = _heading_level(style)
                node = HeadingNode(id=new_node_id(), parent_id=root.id, level=level, style=style)
                node.tags.append(f"H{level}")
            else:
                node = ParagraphNode(
                    id=new_node_id(),
                    parent_id=root.id,
                    style=style or None,
                    alignment=_alignment(para),
                )
            node.reading_order = order
            order += 1

            _parse_inline_children(doc, node, para._p, footnotes)

            root.children.append(node.id)
            doc.add_node(node)

        for table in src.tables:
            order = self._parse_table(doc, root, table, order, footnotes)

        order = _append_footnote_nodes(doc, root, footnotes, order)

        if meta.title:
            doc.accessibility.has_doc_title = True
        return doc

    def _parse_table(
        self,
        doc: CanonicalDocument,
        root,
        table,
        order: int,
        footnotes: dict[str, dict],
    ) -> int:
        tnode = TableNode(
            id=new_node_id(),
            parent_id=root.id,
            rows=len(table.rows),
            cols=len(table.columns),
            reading_order=order,
        )
        for r, row in enumerate(table.rows):
            rownode = TableRowNode(id=new_node_id(), parent_id=tnode.id, row=r)
            for c, cell in enumerate(row.cells):
                cellnode = TableCellNode(id=new_node_id(), parent_id=rownode.id, row=r, col=c)
                for idx, para in enumerate(cell.paragraphs):
                    if idx:
                        sep = RunNode(id=new_node_id(), parent_id=cellnode.id, text="\n")
                        cellnode.children.append(sep.id)
                        doc.add_node(sep)
                    _parse_inline_children(doc, cellnode, para._p, footnotes)
                rownode.children.append(cellnode.id)
                doc.add_node(cellnode)
            tnode.children.append(rownode.id)
            doc.add_node(rownode)
        root.children.append(tnode.id)
        doc.add_node(tnode)
        return order + 1

    def render_preview(self, doc: CanonicalDocument, page: int) -> bytes:
        raise NotImplementedError("DocxAdapter.render_preview — pending LibreOffice/render service")

    def export(self, doc: CanonicalDocument, *, target_mime: str) -> bytes:
        if target_mime != _DOCX_MIME:
            raise NotImplementedError(f"DocxAdapter cannot export to {target_mime}")
        from docos.services.docengine.writers.docx_writer import model_to_docx

        return model_to_docx(doc)


def _heading_level(style: str) -> int:
    try:
        return int(style.replace("Heading", "").strip() or "1")
    except ValueError:
        return 1


def _alignment(para) -> str | None:
    return str(para.alignment) if para.alignment is not None else None


def _parse_footnotes(data: bytes) -> dict[str, dict]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            raw = zf.read("word/footnotes.xml")
    except (KeyError, zipfile.BadZipFile):
        return {}

    root = ET.fromstring(raw)
    notes: dict[str, dict] = {}
    for footnote in root.findall(f"{_W}footnote"):
        fid = footnote.get(f"{_W}id")
        ftype = footnote.get(f"{_W}type")
        if not fid or ftype in {"separator", "continuationSeparator"}:
            continue
        try:
            if int(fid) < 1:
                continue
        except ValueError:
            continue
        marker = str(len(notes) + 1)
        paragraphs: list[list[dict]] = []
        for para in footnote.findall(f"{_W}p"):
            runs = _paragraph_runs(para, include_refs=False)
            if runs:
                paragraphs.append(runs)
        notes[fid] = {
            "footnote_id": fid,
            "marker": marker,
            "paragraphs": paragraphs,
            "attrs": {"ooxml_id": fid},
        }
    return notes


def _parse_inline_children(
    doc: CanonicalDocument,
    parent: ParagraphNode | HeadingNode | TableCellNode,
    paragraph_xml,
    footnotes: dict[str, dict],
) -> None:
    for item in _paragraph_runs(paragraph_xml, include_refs=True):
        if item["kind"] == "footnote_reference":
            fid = item["footnote_id"]
            note = footnotes.get(fid, {})
            ref = FootnoteReferenceNode(
                id=new_node_id(),
                parent_id=parent.id,
                footnote_id=fid,
                marker=str(note.get("marker") or fid),
                attrs={"ooxml_id": fid},
            )
            parent.children.append(ref.id)
            doc.add_node(ref)
            continue
        text = item.get("text", "")
        if not text:
            continue
        rnode = RunNode(
            id=new_node_id(),
            parent_id=parent.id,
            text=text,
            bold=bool(item.get("bold", False)),
            italic=bool(item.get("italic", False)),
            underline=bool(item.get("underline", False)),
            font=item.get("font"),
            size=item.get("size"),
            color=item.get("color"),
        )
        parent.children.append(rnode.id)
        doc.add_node(rnode)


def _append_footnote_nodes(
    doc: CanonicalDocument, root: RootNode, footnotes: dict[str, dict], order: int
) -> int:
    for note in footnotes.values():
        fnode = FootnoteNode(
            id=new_node_id("footnote"),
            parent_id=root.id,
            footnote_id=note["footnote_id"],
            marker=note["marker"],
            reading_order=order,
            attrs=note.get("attrs", {}),
        )
        order += 1
        for para_runs in note["paragraphs"]:
            pnode = ParagraphNode(id=new_node_id(), parent_id=fnode.id)
            for item in para_runs:
                text = item.get("text", "")
                if not text:
                    continue
                rnode = RunNode(
                    id=new_node_id(),
                    parent_id=pnode.id,
                    text=text,
                    bold=bool(item.get("bold", False)),
                    italic=bool(item.get("italic", False)),
                    underline=bool(item.get("underline", False)),
                    font=item.get("font"),
                    size=item.get("size"),
                    color=item.get("color"),
                )
                pnode.children.append(rnode.id)
                doc.add_node(rnode)
            if pnode.children:
                fnode.children.append(pnode.id)
                doc.add_node(pnode)
        root.children.append(fnode.id)
        doc.add_node(fnode)
    return order


def _paragraph_runs(paragraph_xml, *, include_refs: bool) -> list[dict]:
    out: list[dict] = []
    for run in paragraph_xml.findall(f"{_W}r"):
        style = _run_style(run)
        pending: list[str] = []

        for child in list(run):
            if child.tag == f"{_W}t":
                pending.append(child.text or "")
            elif child.tag == f"{_W}tab":
                pending.append("\t")
            elif child.tag == f"{_W}br":
                pending.append("\n")
            elif include_refs and child.tag == f"{_W}footnoteReference":
                if pending:
                    out.append({"kind": "text", "text": "".join(pending), **style})
                    pending.clear()
                fid = child.get(f"{_W}id")
                if fid:
                    out.append({"kind": "footnote_reference", "footnote_id": fid})
        if pending:
            out.append({"kind": "text", "text": "".join(pending), **style})
    return out


def _run_style(run_xml) -> dict:
    rpr = run_xml.find(f"{_W}rPr")
    if rpr is None:
        return {}
    font = None
    rfonts = rpr.find(f"{_W}rFonts")
    if rfonts is not None:
        font = rfonts.get(f"{_W}ascii") or rfonts.get(f"{_W}hAnsi")
    size = None
    sz = rpr.find(f"{_W}sz")
    if sz is not None and sz.get(f"{_W}val"):
        try:
            size = int(sz.get(f"{_W}val")) / 2
        except (TypeError, ValueError):
            size = None
    color = None
    color_el = rpr.find(f"{_W}color")
    if color_el is not None and color_el.get(f"{_W}val") not in (None, "auto"):
        color = f"#{color_el.get(f'{_W}val')}"
    underline = False
    u = rpr.find(f"{_W}u")
    if u is not None and u.get(f"{_W}val") != "none":
        underline = True
    return {
        "bold": rpr.find(f"{_W}b") is not None,
        "italic": rpr.find(f"{_W}i") is not None,
        "underline": underline,
        "font": font,
        "size": size,
        "color": color,
    }
