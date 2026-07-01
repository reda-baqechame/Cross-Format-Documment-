"""Pillar B footnote + future-node compatibility regression tests."""

from __future__ import annotations

import io
import zipfile
from datetime import UTC, datetime
from xml.etree import ElementTree as ET

import fitz
from docx import Document as DocxDocument

from docos.model.document import CanonicalDocument, DocumentMeta
from docos.model.ids import new_doc_id, new_node_id
from docos.model.nodes import (
    FootnoteNode,
    FootnoteReferenceNode,
    ParagraphNode,
    RootNode,
    RunNode,
)
from docos.model.serialize import canonical_hash, from_dict, to_dict
from docos.services.docengine.adapters.docx import _DOCX_MIME, DocxAdapter
from docos.services.docengine.writers.docx_writer import model_to_docx
from docos.services.docengine.writers.markup import model_to_html, model_to_markdown
from docos.services.docengine.writers.searchable_pdf import model_to_searchable_pdf

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_W = f"{{{_W_NS}}}"
_REL = f"{{{_REL_NS}}}"
_CT = f"{{{_CT_NS}}}"

ET.register_namespace("w", _W_NS)


def _doc_with_footnote(secret: str = "Footnote detail") -> CanonicalDocument:
    now = datetime.now(UTC)
    root = RootNode(id=new_node_id("root"))
    doc = CanonicalDocument(
        doc_id=new_doc_id(),
        root_id=root.id,
        meta=DocumentMeta(
            title="Footnote Doc",
            source_format="docx",
            source_mime=_DOCX_MIME,
            created_at=now,
            modified_at=now,
        ),
    )
    doc.add_node(root)
    para = ParagraphNode(id=new_node_id(), parent_id=root.id)
    run = RunNode(id=new_node_id(), parent_id=para.id, text="Body with note")
    ref = FootnoteReferenceNode(
        id=new_node_id(),
        parent_id=para.id,
        footnote_id="fn-1",
        marker="1",
    )
    para.children.extend([run.id, ref.id])
    root.children.append(para.id)
    doc.add_node(para)
    doc.add_node(run)
    doc.add_node(ref)

    footnote = FootnoteNode(
        id=new_node_id("footnote"),
        parent_id=root.id,
        footnote_id="fn-1",
        marker="1",
    )
    note_para = ParagraphNode(id=new_node_id(), parent_id=footnote.id)
    note_run = RunNode(id=new_node_id(), parent_id=note_para.id, text=secret)
    note_para.children.append(note_run.id)
    footnote.children.append(note_para.id)
    root.children.append(footnote.id)
    doc.add_node(footnote)
    doc.add_node(note_para)
    doc.add_node(note_run)
    return doc


def _docx_fixture_with_footnote(
    note_text: str = "Imported footnote", *, in_table: bool = False
) -> bytes:
    doc = DocxDocument()
    if in_table:
        table = doc.add_table(rows=1, cols=1)
        p = table.cell(0, 0).paragraphs[0]
    else:
        p = doc.add_paragraph()
    p.add_run("Body before")
    p.add_run("[[FN]]")
    p.add_run(" after.")
    buf = io.BytesIO()
    doc.save(buf)

    with zipfile.ZipFile(io.BytesIO(buf.getvalue()), "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}

    document = ET.fromstring(entries["word/document.xml"])
    for run in document.findall(f".//{_W}r"):
        text = run.find(f"{_W}t")
        if text is None or text.text != "[[FN]]":
            continue
        for child in list(run):
            run.remove(child)
        ET.SubElement(run, f"{_W}footnoteReference", {f"{_W}id": "2"})
    entries["word/document.xml"] = ET.tostring(document, encoding="utf-8", xml_declaration=True)

    footnotes = ET.Element(f"{_W}footnotes")
    note = ET.SubElement(footnotes, f"{_W}footnote", {f"{_W}id": "2"})
    para = ET.SubElement(note, f"{_W}p")
    run = ET.SubElement(para, f"{_W}r")
    text = ET.SubElement(run, f"{_W}t")
    text.text = note_text
    entries["word/footnotes.xml"] = ET.tostring(footnotes, encoding="utf-8", xml_declaration=True)

    rels = ET.fromstring(entries["word/_rels/document.xml.rels"])
    ET.SubElement(
        rels,
        f"{_REL}Relationship",
        {
            "Id": "rIdFootnotes",
            "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
            "Target": "footnotes.xml",
        },
    )
    entries["word/_rels/document.xml.rels"] = ET.tostring(
        rels, encoding="utf-8", xml_declaration=True
    )

    ctypes = ET.fromstring(entries["[Content_Types].xml"])
    ET.SubElement(
        ctypes,
        f"{_CT}Override",
        {
            "PartName": "/word/footnotes.xml",
            "ContentType": (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
            ),
        },
    )
    entries["[Content_Types].xml"] = ET.tostring(ctypes, encoding="utf-8", xml_declaration=True)

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, payload in entries.items():
            zout.writestr(name, payload)
    return out.getvalue()


def test_footnote_nodes_round_trip_and_hash_is_stable():
    doc = _doc_with_footnote()
    restored = from_dict(to_dict(doc))

    assert any(n.type == "footnote_reference" for n in restored.nodes.values())
    assert any(n.type == "footnote" for n in restored.nodes.values())
    assert canonical_hash(restored) == canonical_hash(from_dict(to_dict(restored)))


def test_unknown_future_node_deserializes_as_unsupported():
    doc = _doc_with_footnote()
    payload = to_dict(doc)
    root_id = payload["root_id"]
    payload["nodes"]["future-1"] = {
        "id": "future-1",
        "type": "equation",
        "parent_id": root_id,
        "children": [],
        "attrs": {"latex": "x^2"},
        "tags": [],
    }
    payload["nodes"][root_id]["children"].append("future-1")

    restored = from_dict(payload)
    node = restored.nodes["future-1"]
    assert node.type == "unsupported"
    assert node.original_type == "equation"
    assert node.raw["attrs"]["latex"] == "x^2"


def test_docx_adapter_parses_body_footnotes():
    doc = DocxAdapter().parse(_docx_fixture_with_footnote("Imported footnote"))
    refs = [n for n in doc.nodes.values() if n.type == "footnote_reference"]
    notes = [n for n in doc.nodes.values() if n.type == "footnote"]

    assert refs and refs[0].marker == "1"
    assert notes and notes[0].footnote_id == "2"
    assert any(n.type == "run" and n.text == "Imported footnote" for n in doc.nodes.values())


def test_docx_adapter_parses_table_cell_footnotes():
    doc = DocxAdapter().parse(_docx_fixture_with_footnote("Table footnote", in_table=True))
    cells = [n for n in doc.nodes.values() if n.type == "table_cell"]
    assert cells and any(doc.nodes[cid].type == "footnote_reference" for cid in cells[0].children)
    assert any(n.type == "run" and n.text == "Table footnote" for n in doc.nodes.values())


def test_docx_writer_emits_real_footnote_part_and_round_trips():
    out = model_to_docx(_doc_with_footnote("Exported footnote"))
    with zipfile.ZipFile(io.BytesIO(out)) as zf:
        document_xml = zf.read("word/document.xml").decode()
        footnotes_xml = zf.read("word/footnotes.xml").decode()
        content_types = zf.read("[Content_Types].xml").decode()
        rels = zf.read("word/_rels/document.xml.rels").decode()

    assert "DOCOS_FOOTNOTE_REF" not in document_xml
    assert "footnoteReference" in document_xml
    assert "Exported footnote" in footnotes_xml
    assert "footnotes+xml" in content_types
    assert "relationships/footnotes" in rels

    reparsed = DocxAdapter().parse(out)
    assert any(n.type == "run" and n.text == "Exported footnote" for n in reparsed.nodes.values())
    assert b"Footnotes" in model_to_markdown(_doc_with_footnote("Exported footnote"))
    assert b"Footnotes" in model_to_html(_doc_with_footnote("Exported footnote"))


def test_redacted_footnote_text_is_removed_from_exports():
    secret = "SECRET_FOOTNOTE_TOKEN"
    doc = _doc_with_footnote(secret)
    secret_node = next(n for n in doc.nodes.values() if n.type == "run" and n.text == secret)
    doc.redaction.redacted_node_ids.append(secret_node.id)

    docx = model_to_docx(doc)
    markdown = model_to_markdown(doc)
    html = model_to_html(doc)
    pdf = model_to_searchable_pdf(doc)

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        assert secret not in zf.read("word/footnotes.xml").decode()
    assert secret.encode() not in markdown
    assert secret.encode() not in html
    pdf_doc = fitz.open(stream=pdf, filetype="pdf")
    try:
        text = "\n".join(page.get_text() for page in pdf_doc)
    finally:
        pdf_doc.close()
    assert secret not in text


def test_upload_docx_footnote_model_and_exports(client):
    up = client.post(
        "/documents",
        files={
            "file": (
                "footnote.docx",
                io.BytesIO(_docx_fixture_with_footnote("API footnote")),
                _DOCX_MIME,
            )
        },
    )
    assert up.status_code == 200, up.text
    doc_id = up.json()["doc_id"]

    model = client.get(f"/documents/{doc_id}/model").json()["document"]
    assert any(n["type"] == "footnote_reference" for n in model["nodes"].values())
    assert any(n["type"] == "footnote" for n in model["nodes"].values())

    docx = client.get(f"/documents/{doc_id}/export", params={"format": "docx"})
    html = client.get(f"/documents/{doc_id}/export", params={"format": "html"})
    pdf = client.get(f"/documents/{doc_id}/searchable-pdf")
    assert docx.status_code == 200
    assert b"API footnote" in zipfile.ZipFile(io.BytesIO(docx.content)).read("word/footnotes.xml")
    assert html.status_code == 200 and b"API footnote" in html.content
    assert pdf.status_code == 200
