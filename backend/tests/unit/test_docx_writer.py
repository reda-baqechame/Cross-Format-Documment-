"""The universal DOCX writer rebuilds a real .docx from the canonical model."""

from __future__ import annotations

import io
from datetime import UTC, datetime

from docx import Document as DocxDocument

from docos.model.document import CanonicalDocument, DocumentMeta
from docos.model.ids import new_doc_id, new_node_id
from docos.model.nodes import ParagraphNode, RootNode, RunNode
from docos.services.docengine.adapters.docx import _DOCX_MIME, DocxAdapter
from docos.services.docengine.adapters.txt import TxtAdapter
from docos.services.docengine.writers.docx_writer import model_to_docx


def test_model_to_docx_preserves_heading_runs_and_table(sample_docx_bytes):
    doc = DocxAdapter().parse(sample_docx_bytes)
    out = model_to_docx(doc)

    reopened = DocxDocument(io.BytesIO(out))
    headings = [p for p in reopened.paragraphs if p.style and p.style.name.startswith("Heading")]
    assert any("A Heading" in p.text for p in headings)
    assert any("normal paragraph" in p.text for p in reopened.paragraphs)

    assert reopened.tables, "table should be rebuilt"
    cell_texts = {c.text for t in reopened.tables for row in t.rows for c in row.cells}
    assert "r0c0" in cell_texts and "r1c1" in cell_texts


def test_model_to_docx_keeps_run_formatting():
    now = datetime.now(UTC)
    root = RootNode(id=new_node_id("root"))
    doc = CanonicalDocument(
        doc_id=new_doc_id(),
        root_id=root.id,
        meta=DocumentMeta(
            source_format="txt",
            source_mime="text/plain",
            created_at=now,
            modified_at=now,
        ),
    )
    doc.add_node(root)
    para = ParagraphNode(id=new_node_id(), parent_id=root.id)
    run = RunNode(id=new_node_id(), parent_id=para.id, text="bold words", bold=True, italic=True)
    para.children.append(run.id)
    root.children.append(para.id)
    doc.add_node(para)
    doc.add_node(run)

    reopened = DocxDocument(io.BytesIO(model_to_docx(doc)))
    runs = [r for p in reopened.paragraphs for r in p.runs if r.text]
    assert runs and runs[0].bold and runs[0].italic


def test_docx_adapter_export_roundtrip_keeps_run_text(sample_docx_bytes):
    doc = DocxAdapter().parse(sample_docx_bytes)
    exported = DocxAdapter().export(doc, target_mime=_DOCX_MIME)
    reparsed = DocxAdapter().parse(exported)
    original = {n.text for n in doc.nodes.values() if n.type == "run" and n.text}
    after = {n.text for n in reparsed.nodes.values() if n.type == "run" and n.text}
    assert original.issubset(after)


def test_txt_origin_exports_as_docx():
    doc = TxtAdapter().parse(b"First block\n\nSecond block")
    out = model_to_docx(doc)
    reopened = DocxDocument(io.BytesIO(out))
    text = "\n".join(p.text for p in reopened.paragraphs)
    assert "First block" in text and "Second block" in text
